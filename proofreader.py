#!/usr/bin/env python3
"""
Document Proofreader class - handles document processing and OpenAI integration
"""

import os
import json
import re
from docx import Document
import tiktoken
from openai import AzureOpenAI

from config import Config
from utils import calculate_costs


class DocumentProofreader:
    def __init__(self):
        """Initialize the document proofreader with Azure OpenAI client"""
        # Validate configuration first
        Config.validate()
        
        self.client = AzureOpenAI(
            api_key=Config.AZURE_OPENAI_API_KEY,
            api_version=Config.AZURE_API_VERSION,
            azure_endpoint=Config.AZURE_OPENAI_ENDPOINT
        )
        self.encoder = tiktoken.encoding_for_model("gpt-4")
    
    def extract_tables(self, doc):
        """
        Extract all tables from the document with location context
        Note: This method stores basic table info. Enhanced location context is added later in extract_document_structure.
        """
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:  # Only include non-empty cells
                        row_data.append({
                            'text': cell_text,
                            'location': f"Table {table_idx + 1}, Row {row_idx + 1}, Column {col_idx + 1}",
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'col_idx': col_idx
                        })
                if row_data:  # Only include rows with content
                    table_data.append(row_data)
            
            if table_data:  # Only include tables with content
                tables.append({
                    'table_id': table_idx,
                    'location': f"Table {table_idx + 1}",
                    'rows': table_data,
                    'content_text': self._table_to_text(table_data, table_idx),
                    'token_count': len(self.encoder.encode(self._table_to_text(table_data, table_idx)))
                })
        
        return tables
    
    def _table_to_text(self, table_data, table_idx=0):
        """Convert table data to readable text format for proofreading"""
        text_lines = []
        for row_idx, row in enumerate(table_data):
            row_text = " | ".join([cell['text'] for cell in row])
            text_lines.append(f"Row {row_idx + 1}: {row_text}")
        return "\n".join(text_lines)
    
    def detect_list_paragraphs(self, doc):
        """
        Detect paragraphs that are part of lists (bulleted or numbered)
        """
        list_paragraphs = []
        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
                
            # Check if paragraph has list formatting
            if para._element.pPr is not None:
                numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    # This is a numbered list item
                    list_paragraphs.append({
                        'paragraph_index': para_idx,
                        'text': para.text,
                        'list_type': 'numbered',
                        'location': f"Paragraph {para_idx + 1} (Numbered list)",
                        'token_count': len(self.encoder.encode(para.text))
                    })
                    continue
            
            # Check for bullet point patterns in text
            text = para.text.strip()
            if (text.startswith('•') or text.startswith('-') or text.startswith('*') or
                (len(text) > 2 and text[1] == '.' and text[0].isdigit())):
                list_paragraphs.append({
                    'paragraph_index': para_idx,
                    'text': text,
                    'list_type': 'bulleted' if not text[0].isdigit() else 'numbered',
                    'location': f"Paragraph {para_idx + 1} (List item)",
                    'token_count': len(self.encoder.encode(text))
                })
        
        return list_paragraphs
    
    def _build_section_path(self, section, sections):
        """Build hierarchical path to section using heading titles"""
        path_parts = []
        current = section
        
        # Build path by traversing up the hierarchy
        while current and current.get('title') != 'Document Root':
            # Only include meaningful section titles (level 1 and 2 headings)
            if current.get('level', 0) <= 2 and current.get('title'):
                path_parts.append(current['title'])
            
            # Move to parent section
            parent_idx = current.get('parent_idx')
            if parent_idx is not None and parent_idx < len(sections):
                current = sections[parent_idx]
            else:
                break
        
        # Reverse to get correct order (parent → child)
        path_parts.reverse()
        return ' → '.join(path_parts) if path_parts else ""
    
    def _enhance_table_content_locations(self, content_text, table_rows, section_context):
        """Enhance table content with section context in location markers"""
        enhanced_lines = []
        base_section = section_context.split(' → ')[0] if ' → ' in section_context else ""
        
        for row_idx, row in enumerate(table_rows):
            row_cells = []
            for cell in row:
                cell_text = cell['text']
                basic_location = f"Table {cell['table_idx'] + 1}, Row {cell['row_idx'] + 1}, Column {cell['col_idx'] + 1}"
                if base_section:
                    enhanced_location = f"{base_section} → {basic_location}"
                else:
                    enhanced_location = basic_location
                row_cells.append(f"[{enhanced_location}] {cell_text}")
            
            if row_cells:
                enhanced_lines.append(f"Row {row_idx + 1}: " + " | ".join(row_cells))
        
        return "\n".join(enhanced_lines)
        
    def extract_document_structure(self, docx_path):
        """
        Extract document structure using python-docx
        Returns a list of sections with hierarchical information, plus tables and lists
        """
        print(f"Opening document: {docx_path}")
        doc = Document(docx_path)
        print(f"Document opened successfully with {len(doc.paragraphs)} paragraphs")
        
        # Extract tables and lists first
        tables = self.extract_tables(doc)
        list_paragraphs = self.detect_list_paragraphs(doc)
        
        # Create lookup for list paragraphs
        list_para_indices = {item['paragraph_index'] for item in list_paragraphs}
        
        print(f"Found {len(tables)} tables and {len(list_paragraphs)} list items")
        
        sections = []
        
        # Start with a default root section
        root_section = {
            "title": "Document Root",
            "content": [],
            "level": 0,
            "parent_idx": None,
            "section_id": 0,
            "children": [],
            "token_count": 0
        }
        sections.append(root_section)
        
        current_section = root_section
        section_stack = [root_section]  # Track section hierarchy
        
        # Function to determine heading level from style name
        def get_heading_level(paragraph):
            if not paragraph.style.name.startswith('Heading'):
                return None
            try:
                return int(paragraph.style.name.replace('Heading ', ''))
            except ValueError:
                return 1  # Default to level 1 if can't parse

        # Process all paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                continue
                
            heading_level = get_heading_level(para)
            
            # Add location information to paragraph text with section context
            para_location = f"Paragraph {para_idx + 1}"
            if para_idx in list_para_indices:
                list_item = next(item for item in list_paragraphs if item['paragraph_index'] == para_idx)
                para_location = list_item['location']
            
            # Add section context to location
            section_path = self._build_section_path(current_section, sections)
            if section_path and section_path != "Document Root":
                para_location = f"{section_path} → {para_location}"
            
            if heading_level is not None:
                # This is a heading - create a new section
                
                # Adjust the section stack based on heading level
                while len(section_stack) > 1 and section_stack[-1]["level"] >= heading_level:
                    section_stack.pop()
                
                # Find parent in the section stack
                parent = section_stack[-1]  # Get the last item in the stack
                parent_idx = parent["section_id"]
                
                # Create new section
                new_section = {
                    "title": para.text,
                    "content": [],
                    "level": heading_level,
                    "parent_idx": parent_idx,
                    "section_id": len(sections),
                    "children": [],
                    "token_count": 0
                }
                
                # Add this section as a child to its parent
                if parent_idx < len(sections):  # Ensure parent index is valid
                    sections[parent_idx]["children"].append(len(sections))
                
                # Add to sections list
                sections.append(new_section)
                
                # Update current section and stack
                current_section = new_section
                section_stack.append(new_section)
            else:
                # Regular paragraph - add to current section with location info
                current_section["content"].append({
                    'text': para.text,
                    'location': para_location,
                    'paragraph_index': para_idx
                })
        
        # Add tables as separate sections with enhanced location context
        for table in tables:
            # Find the section context where this table appears (approximate by finding the last heading section)
            table_parent_section = None
            for section in reversed(sections):
                if section.get('level', 0) > 0:  # Find the most recent heading section
                    table_parent_section = section
                    break
            
            # Build enhanced table location
            table_location = table['location']
            if table_parent_section:
                section_path = self._build_section_path(table_parent_section, sections)
                if section_path:
                    table_location = f"{section_path} → {table['location']}"
            
            # Update table content with enhanced location context
            enhanced_content = self._enhance_table_content_locations(table['content_text'], table['rows'], table_location)
            
            table_section = {
                "title": table_location,
                "content": [{'text': enhanced_content, 'location': table_location, 'type': 'table'}],
                "level": 0,
                "parent_idx": table_parent_section['section_id'] if table_parent_section else None,
                "section_id": len(sections),
                "children": [],
                "token_count": len(self.encoder.encode(enhanced_content))
            }
            sections.append(table_section)
        
        # Calculate token counts for all sections
        for section in sections:
            if section.get("token_count", 0) == 0:  # Only calculate if not already set (for tables)
                section_texts = []
                for content_item in section["content"]:
                    if isinstance(content_item, dict):
                        section_texts.append(content_item['text'])
                    else:
                        section_texts.append(content_item)  # Handle old format
                section_text = "\n".join(section_texts)
                section["token_count"] = len(self.encoder.encode(section_text))
        
        # Log section information
        print(f"Extracted {len(sections)} sections:")
        for i, section in enumerate(sections):
            print(f"  Section {i}: '{section['title']}' - Level {section['level']} - {len(section['content'])} paragraphs")
        
        return sections
    
    def segment_document(self, sections, max_tokens=None):
        """Segment the document based on structural boundaries and token limits"""
        if max_tokens is None:
            max_tokens = Config.MAX_TOKENS_PER_SEGMENT
            
        # Skip the root section (if it exists)
        content_sections = sections[1:] if len(sections) > 1 else sections
        
        # If we have very few sections, just return them all as one segment
        if len(content_sections) <= 1:
            return [[0]] if len(sections) == 1 else [[0, 1]]
        
        # Extract section texts and calculate token counts
        token_counts = []
        
        for section in content_sections:
            token_counts.append(section["token_count"])
        
        # Handle empty sections
        if all(count == 0 for count in token_counts):
            print("Warning: No content found in sections. Returning all as one segment.")
            return [list(range(len(sections)))]
        
        print(f"Segmenting {len(content_sections)} content sections")
        
        # Create segments based on document structure and token limits
        segments = []
        current_segment = []
        current_tokens = 0
        
        # Start with the root section if it exists
        if len(sections) > 1:
            current_segment = [0]
        
        # Process each content section
        for i, section in enumerate(content_sections):
            # The actual section index in the full sections list
            section_idx = i + 1 if len(sections) > 1 else i
            section_token_count = token_counts[i]
            
            # If this would be the first section in the segment or we're under token limit
            if not current_segment or (current_tokens + section_token_count <= max_tokens):
                # Add to current segment
                current_segment.append(section_idx)
                current_tokens += section_token_count
            else:
                # This section would exceed the token limit
                # Save current segment and start a new one
                if current_segment:
                    segments.append(current_segment)
                current_segment = [section_idx]
                current_tokens = section_token_count
                
            # Special case: if this section alone exceeds the token limit,
            # we need to split it further in the future (just add warning for now)
            if section_token_count > max_tokens:
                print(f"Warning: Section {section_idx} ('{section['title']}') exceeds maximum token limit ({section_token_count} > {max_tokens})")
        
        # Add the last segment
        if current_segment:
            segments.append(current_segment)
            
        print(f"Created {len(segments)} segments")
        for i, segment in enumerate(segments):
            segment_token_count = sum(sections[idx]["token_count"] for idx in segment)
            print(f"  Segment {i}: {len(segment)} sections - {segment_token_count} tokens - {[sections[idx]['title'] for idx in segment]}")
            
        return segments
    
    def proofread_content(self, content):
        """Send content to GPT-4o for proofreading"""
        try:
            user_prompt = Config.USER_PROMPT_TEMPLATE.format(content=content)
            
            response = self.client.chat.completions.create(
                model=Config.MODEL_NAME,
                messages=[
                    {"role": "system", "content": Config.SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"Error proofreading content: {str(e)}")
            return f"Error during proofreading: {str(e)}"
    
    def create_docx_report(self, content, output_path, include_costs=False, input_tokens=0, output_tokens=0):
        """Create a nicely formatted DOCX report from markdown content"""
        doc = Document()
        
        # Add document title
        doc.add_heading('Document Proofreading Report', 0)
        
        # Add cost analysis if requested
        if include_costs:
            costs = calculate_costs(input_tokens, output_tokens)
            doc.add_heading('Cost Analysis', 1)
            doc.add_paragraph(f"• Total Input Tokens: {costs['input_tokens']:,}")
            doc.add_paragraph(f"• Total Output Tokens: {costs['output_tokens']:,}")
            doc.add_paragraph(f"• Estimated API Cost: ${costs['total_cost']:.4f}")
        
        # Parse markdown content and convert to docx
        lines = content.split('\n')
        in_code_block = False
        
        for line in lines:
            # Handle headings (## Heading)
            if line.startswith('#'):
                level = len(line.split(' ')[0])  # Count the # symbols
                heading_text = line.lstrip('#').strip()
                doc.add_heading(heading_text, level)
                continue
                
            # Handle issue blocks
            if line.startswith('### Issue'):
                doc.add_heading(line.strip('# '), 3)
                continue
                
            # Handle bold text
            if '**Original**:' in line:
                p = doc.add_paragraph()
                p.add_run('Original: ').bold = True
                p.add_run(line.split('**Original**:')[1].strip())
                continue
                
            if '**Issue**:' in line:
                p = doc.add_paragraph()
                p.add_run('Issue: ').bold = True
                p.add_run(line.split('**Issue**:')[1].strip())
                continue
                
            if '**Suggestion**:' in line:
                p = doc.add_paragraph()
                p.add_run('Suggestion: ').bold = True
                p.add_run(line.split('**Suggestion**:')[1].strip())
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
                
            # Handle separators
            if line.startswith('---'):
                doc.add_paragraph('').add_run('─' * 50)
                continue
                
            # Handle regular text
            if line.strip() and not in_code_block:
                doc.add_paragraph(line)
        
        # Save the document
        doc.save(output_path)
        return output_path
    
    def create_html_report(self, content, output_path, include_costs=False, input_tokens=0, output_tokens=0):
        """Create a nicely formatted HTML report from markdown content"""
        from utils import calculate_costs
        
        html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Proofreading Report</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f8f9fa;
        }}
        .container {{
            background: white;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            padding: 30px;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        h3 {{
            color: #e74c3c;
            background: #fdf2f2;
            padding: 10px;
            border-left: 4px solid #e74c3c;
            margin: 20px 0 10px 0;
        }}
        .cost-analysis {{
            background: #e8f5e8;
            border: 1px solid #4caf50;
            border-radius: 5px;
            padding: 15px;
            margin: 20px 0;
        }}
        .cost-analysis h2 {{
            color: #2e7d32;
            margin-top: 0;
        }}
        .issue-block {{
            background: #fff;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            margin: 15px 0;
            padding: 20px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        .location {{
            background: #3498db;
            color: white;
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 0.9em;
            font-weight: bold;
        }}
        .comparison-table {{
            width: 100%;
            margin: 15px 0;
            border-collapse: collapse;
            font-family: monospace;
        }}
        .comparison-table th {{
            background: #f8f9fa;
            padding: 12px;
            text-align: left;
            font-weight: bold;
            border: 1px solid #dee2e6;
        }}
        .comparison-table .original-header {{
            color: #c62828;
            border-left: 4px solid #f44336;
        }}
        .comparison-table .suggestion-header {{
            color: #2e7d32;
            border-left: 4px solid #4caf50;
        }}
        .comparison-table td {{
            padding: 15px;
            border: 1px solid #dee2e6;
            vertical-align: top;
        }}
        .comparison-table .original-cell {{
            background: #ffebee;
            border-left: 4px solid #f44336;
        }}
        .comparison-table .suggestion-cell {{
            background: #e8f5e8;
            border-left: 4px solid #4caf50;
        }}
        .issue-checkbox {{
            float: right;
            margin-left: 15px;
        }}
        .issue-checkbox input[type="checkbox"] {{
            transform: scale(1.2);
            margin-right: 8px;
        }}
        .issue-checkbox label {{
            font-size: 0.9em;
            color: #666;
            cursor: pointer;
        }}
        .completed {{
            opacity: 0.7;
            background: #f8f9fa !important;
        }}
        .completed .issue-type {{
            opacity: 0.8;
        }}
        .progress-bar {{
            background: #e9ecef;
            border-radius: 4px;
            height: 8px;
            margin: 15px 0;
            overflow: hidden;
        }}
        .progress-fill {{
            background: #28a745;
            height: 100%;
            width: 0%;
            transition: width 0.3s ease;
        }}
        .progress-text {{
            text-align: center;
            font-size: 0.9em;
            color: #666;
            margin-bottom: 10px;
        }}
        .issue-type {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 0.8em;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .spelling {{ background: #ffebee; color: #c62828; }}
        .grammar {{ background: #fff3e0; color: #ef6c00; }}
        .consistency {{ background: #e3f2fd; color: #1565c0; }}
        .formatting {{ background: #f3e5f5; color: #7b1fa2; }}
        .clarity {{ background: #e0f2f1; color: #00695c; }}
        .summary {{
            background: #f5f5f5;
            border-radius: 8px;
            padding: 20px;
            margin: 20px 0;
        }}
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
            .issue-checkbox {{ display: none; }}
            .progress-bar, .progress-text {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Document Proofreading Report</h1>
        {cost_section}
        <div class="progress-text">Progress: <span id="progress-count">0</span> of <span id="total-count">0</span> issues reviewed</div>
        <div class="progress-bar">
            <div class="progress-fill" id="progress-fill"></div>
        </div>
        {content_html}
    </div>
    
    <script>
        // Initialize progress tracking
        document.addEventListener('DOMContentLoaded', function() {{
            const checkboxes = document.querySelectorAll('.issue-checkbox input[type="checkbox"]');
            const totalCount = checkboxes.length;
            const progressCount = document.getElementById('progress-count');
            const totalCountSpan = document.getElementById('total-count');
            const progressFill = document.getElementById('progress-fill');
            
            totalCountSpan.textContent = totalCount;
            
            // Load saved states from localStorage
            checkboxes.forEach(function(checkbox, index) {{
                const issueId = 'issue_' + index;
                checkbox.id = issueId;
                
                const savedState = localStorage.getItem(issueId);
                if (savedState === 'true') {{
                    checkbox.checked = true;
                    checkbox.closest('.issue-block').classList.add('completed');
                }}
                
                // Add event listener for state changes
                checkbox.addEventListener('change', function() {{
                    const issueBlock = this.closest('.issue-block');
                    
                    if (this.checked) {{
                        issueBlock.classList.add('completed');
                        localStorage.setItem(issueId, 'true');
                    }} else {{
                        issueBlock.classList.remove('completed');
                        localStorage.setItem(issueId, 'false');
                    }}
                    
                    updateProgress();
                }});
            }});
            
            function updateProgress() {{
                const checkedCount = document.querySelectorAll('.issue-checkbox input[type="checkbox"]:checked').length;
                const percentage = totalCount > 0 ? (checkedCount / totalCount) * 100 : 0;
                
                progressCount.textContent = checkedCount;
                progressFill.style.width = percentage + '%';
            }}
            
            // Initial progress update
            updateProgress();
        }});
    </script>
</body>
</html>
"""
        
        cost_section = ""
        if include_costs:
            costs = calculate_costs(input_tokens, output_tokens)
            cost_section = f"""
        <div class="cost-analysis">
            <h2>Cost Analysis</h2>
            <ul>
                <li><strong>Total Input Tokens:</strong> {costs['input_tokens']:,}</li>
                <li><strong>Total Output Tokens:</strong> {costs['output_tokens']:,}</li>
                <li><strong>Estimated API Cost:</strong> ${costs['total_cost']:.4f}</li>
            </ul>
        </div>"""
        
        # Convert markdown content to HTML
        content_html = self._markdown_to_html(content)
        
        # Fill template
        html_content = html_template.format(
            cost_section=cost_section,
            content_html=content_html
        )
        
        # Save HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def _markdown_to_html(self, markdown_content):
        """Convert markdown content to HTML with custom styling"""
        lines = markdown_content.split('\n')
        html_lines = []
        in_issue_block = False
        current_issue = {}
        
        for line in lines:
            line = line.strip()
            
            if line.startswith('### Issue'):
                # Start new issue block
                if in_issue_block and current_issue:
                    html_lines.append(self._format_issue_block(current_issue))
                
                in_issue_block = True
                # Extract issue type
                if ':' in line:
                    issue_type = line.split(':')[1].strip().replace('[', '').replace(']', '')
                    current_issue = {'type': issue_type, 'details': {}}
                else:
                    current_issue = {'type': 'General', 'details': {}}
                continue
                
            elif line.startswith('**Location**:') and in_issue_block:
                current_issue['details']['location'] = line.replace('**Location**:', '').strip()
                continue
                
            elif line.startswith('**Original**:') and in_issue_block:
                current_issue['details']['original'] = line.replace('**Original**:', '').strip().strip('"')
                continue
                
            elif line.startswith('**Issue**:') and in_issue_block:
                current_issue['details']['issue'] = line.replace('**Issue**:', '').strip()
                continue
                
            elif line.startswith('**Suggestion**:') and in_issue_block:
                current_issue['details']['suggestion'] = line.replace('**Suggestion**:', '').strip().strip('"')
                continue
                
            elif line.startswith('##') and not line.startswith('###'):
                # End issue block if we hit a new section
                if in_issue_block and current_issue:
                    html_lines.append(self._format_issue_block(current_issue))
                    in_issue_block = False
                    current_issue = {}
                
                html_lines.append(f'<h2>{line.replace("##", "").strip()}</h2>')
                continue
                
            elif line.startswith('---'):
                # End issue block at separator
                if in_issue_block and current_issue:
                    html_lines.append(self._format_issue_block(current_issue))
                    in_issue_block = False
                    current_issue = {}
                continue
                
            elif line and not in_issue_block:
                # Regular content
                if line.startswith('Total'):
                    html_lines.append(f'<div class="summary"><p>{line}</p></div>')
                else:
                    html_lines.append(f'<p>{line}</p>')
        
        # Handle last issue block
        if in_issue_block and current_issue:
            html_lines.append(self._format_issue_block(current_issue))
        
        return '\n'.join(html_lines)
    
    def _format_issue_block(self, issue_data):
        """Format an individual issue as HTML with side-by-side comparison"""
        issue_type = issue_data.get('type', 'General').lower()
        details = issue_data.get('details', {})
        
        # Determine CSS class for issue type
        css_class = 'spelling'
        if 'grammar' in issue_type:
            css_class = 'grammar'
        elif 'consistency' in issue_type:
            css_class = 'consistency'
        elif 'formatting' in issue_type:
            css_class = 'formatting'
        elif 'clarity' in issue_type:
            css_class = 'clarity'
        
        html = f'<div class="issue-block">'
        
        # Header with issue type, location and checkbox
        html += f'<span class="issue-type {css_class}">{issue_data.get("type", "General")}</span>'
        
        # Add checkbox for tracking completion
        html += f'''<div class="issue-checkbox">
            <input type="checkbox" id="issue-checkbox">
            <label for="issue-checkbox">Reviewed</label>
        </div>'''
        
        if 'location' in details:
            html += f'<br><span class="location">{details["location"]}</span>'
        
        # Issue description on top
        if 'issue' in details:
            html += f'<p><strong>Issue:</strong> {details["issue"]}</p>'
        
        # Side-by-side comparison table
        if 'original' in details and 'suggestion' in details:
            html += f'''
            <table class="comparison-table">
                <thead>
                    <tr>
                        <th class="original-header">Original Text</th>
                        <th class="suggestion-header">Suggested Text</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td class="original-cell">{details["original"]}</td>
                        <td class="suggestion-cell">{details["suggestion"]}</td>
                    </tr>
                </tbody>
            </table>'''
        elif 'original' in details:
            html += f'<div class="original-cell" style="padding: 15px; margin: 10px 0;"><strong>Original:</strong> {details["original"]}</div>'
        elif 'suggestion' in details:
            html += f'<div class="suggestion-cell" style="padding: 15px; margin: 10px 0;"><strong>Suggestion:</strong> {details["suggestion"]}</div>'
        
        html += '</div>'
        return html
    
    def combine_results(self, results, output_path, docx_output_path, total_input_tokens=0, total_output_tokens=0):
        """Combine all segment results into a single document"""
        # Create a combined markdown document
        with open(output_path, "w") as f:
            f.write("# Complete Document Proofreading Report\n\n")
            f.write("## Overview\n\n")
            f.write(f"Total segments processed: {len(results)}\n\n")
            
            # Calculate statistics
            total_issues = 0
            issue_types = {}
            
            for i, result in enumerate(results):
                feedback = result["feedback"]
                # Try to count issues using regex
                issue_matches = re.findall(r'### Issue \d+:', feedback)
                segment_issues = len(issue_matches)
                total_issues += segment_issues
                
                # Try to extract issue types
                type_matches = re.findall(r'### Issue \d+: \[(.*?)\]', feedback)
                for issue_type in type_matches:
                    issue_types[issue_type] = issue_types.get(issue_type, 0) + 1
            
            # Add cost calculation
            costs = calculate_costs(total_input_tokens, total_output_tokens)
            
            # Write cost information
            f.write("### Cost Analysis\n")
            f.write(f"- Total Input Tokens: {costs['input_tokens']:,}\n")
            f.write(f"- Total Output Tokens: {costs['output_tokens']:,}\n")
            f.write(f"- Estimated API Cost: ${costs['total_cost']:.4f}\n\n")
            
            # Write statistics
            f.write(f"Total issues identified: {total_issues}\n\n")
            
            if issue_types:
                f.write("### Issue Types\n")
                for issue_type, count in sorted(issue_types.items(), key=lambda x: x[1], reverse=True):
                    f.write(f"- {issue_type}: {count}\n")
                f.write("\n")
            
            # Write each segment's feedback
            for i, result in enumerate(results):
                f.write(f"## Segment {i+1}\n\n")
                f.write("### Sections Included\n")
                for title in result["section_titles"]:
                    f.write(f"- {title}\n")
                f.write("\n### Feedback\n\n")
                f.write(result["feedback"])
                f.write("\n\n---\n\n")
        
        # Create DOCX version (simplified)
        self.create_docx_report(
            f"# Complete Document Proofreading Report\n\nTotal segments: {len(results)}\nTotal issues: {total_issues}\n\n" +
            "\n\n".join([result["feedback"] for result in results]),
            docx_output_path,
            include_costs=True,
            input_tokens=total_input_tokens,
            output_tokens=total_output_tokens
        )
        
        # Create HTML version
        html_output_path = docx_output_path.replace('.docx', '.html')
        self.create_html_report(
            f"# Complete Document Proofreading Report\n\nTotal segments: {len(results)}\nTotal issues: {total_issues}\n\n" +
            "\n\n".join([result["feedback"] for result in results]),
            html_output_path,
            include_costs=True,
            input_tokens=total_input_tokens,
            output_tokens=total_output_tokens
        )
        
        # Also save JSON format for programmatic use
        combined_json = {
            "segments": results,
            "statistics": {
                "total_segments": len(results),
                "total_issues": total_issues,
                "issue_types": issue_types,
                "tokens": {
                    "input": total_input_tokens,
                    "output": total_output_tokens
                },
                "costs": costs
            },
            "docx_path": docx_output_path,
            "html_path": html_output_path
        }
        
        json_path = output_path.replace(".md", ".json")
        with open(json_path, "w") as f:
            json.dump(combined_json, f, indent=2)
        
        return combined_json