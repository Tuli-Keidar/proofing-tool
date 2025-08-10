#!/usr/bin/env python3
"""
Main entry point for the Document Proofreader tool
"""

import os
import sys
import json
import argparse
from docx import Document

from config import Config
from proofreader import DocumentProofreader
from utils import calculate_costs


def create_test_document():
    """Create a test document with intentional errors for testing"""
    try:
        print("Creating a test document as fallback...")
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph with a typo and misspelling. There are grammer errors here and incorect words.')
        doc.add_heading('Section Two', level=1)
        doc.add_paragraph('This is a second paragraph with more mistakes and erors. Its important to fix all these isues.')
        doc.save('test_document.docx')
        print(f"Created test document: test_document.docx")
        return 'test_document.docx'
    except Exception as e:
        print(f"Failed to create test document: {str(e)}")
        return None


def proofread_document(document_path=None, output_dir=None):
    """
    Main proofreading function
    
    Parameters:
    -----------
    document_path : str, optional
        Path to the DOCX document to proofread
    output_dir : str, optional
        Directory to save the results
        
    Returns:
    --------
    dict
        Results from the proofreading process
    """
    # Use defaults from config if not provided
    if document_path is None:
        document_path = Config.INPUT_DOCUMENT_PATH
    if output_dir is None:
        output_dir = Config.OUTPUT_DIRECTORY
    
    print(f"Starting proofreading of: {document_path}")
    
    # Verify the document exists
    if not os.path.exists(document_path):
        print(f"ERROR: Document not found at path: {document_path}")
        print(f"Current working directory: {os.getcwd()}")
        print("Available .docx files:")
        docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
        if docx_files:
            for f in docx_files:
                print(f"  - {f}")
        else:
            print("  No .docx files found")
        
        # Try to create a test document as a fallback
        test_doc = create_test_document()
        if test_doc:
            document_path = test_doc
        else:
            return {"error": "Document not found and could not create test document"}
    
    # Create output directory
    os.makedirs(output_dir, exist_ok=True)
    
    # Initialize the proofreader
    try:
        proofreader = DocumentProofreader()
        print("✓ Document proofreader initialized successfully")
    except Exception as e:
        print(f"✗ Failed to initialize proofreader: {e}")
        return {"error": f"Failed to initialize proofreader: {e}"}
    
    print(f"Extracting document structure from {document_path}...")
    
    try:
        # Extract document structure
        sections = proofreader.extract_document_structure(document_path)
        
        # If we have no sections with content, extract full document as plain text
        if all(len(section["content"]) == 0 for section in sections):
            print("No structured content found. Reading document as plain text.")
            doc = Document(document_path)
            all_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Proofread the entire document as one segment
            print("Proofreading full document content...")
            feedback = proofreader.proofread_content(all_text)
            
            # Estimate tokens
            input_tokens = len(proofreader.encoder.encode(all_text))
            output_tokens = len(proofreader.encoder.encode(feedback))
            
            # Save results
            with open(os.path.join(output_dir, "proofreading_results.md"), "w") as f:
                f.write("# Document Proofreading Results\n\n")
                
                # Add cost calculations
                costs = calculate_costs(input_tokens, output_tokens)
                f.write("## Cost Analysis\n")
                f.write(f"- Total Input Tokens: {costs['input_tokens']:,}\n")
                f.write(f"- Total Output Tokens: {costs['output_tokens']:,}\n")
                f.write(f"- Estimated API Cost: ${costs['total_cost']:.4f}\n\n")
                
                f.write("## Feedback\n\n")
                f.write(feedback)
            
            # Create DOCX output
            docx_output_path = os.path.join(output_dir, "proofreading_results.docx")
            proofreader.create_docx_report(feedback, docx_output_path, 
                                          include_costs=True, input_tokens=input_tokens, output_tokens=output_tokens)
            
            print(f"Proofreading complete. Results saved to {output_dir}/proofreading_results.md")
            print(f"DOCX report saved to {docx_output_path}")
            
            return {
                "document": document_path,
                "feedback": feedback,
                "docx_report": docx_output_path,
                "tokens": {
                    "input": input_tokens,
                    "output": output_tokens
                },
                "costs": calculate_costs(input_tokens, output_tokens)
            }
            
        # Segment the document
        print("Segmenting document...")
        segments = proofreader.segment_document(sections)
        
        # Process each segment
        results = []
        total_input_tokens = 0
        total_output_tokens = 0
        
        for i, segment_indices in enumerate(segments):
            segment_content = []
            segment_input_tokens = 0
            
            for idx in segment_indices:
                section = sections[idx]
                segment_content.append(f"## {section['title']}")
                segment_content.extend(section["content"])
                segment_input_tokens += section["token_count"]
            
            content = "\n\n".join(segment_content)
            total_input_tokens += segment_input_tokens
            
            print(f"Proofreading segment {i+1}/{len(segments)}...")
            feedback = proofreader.proofread_content(content)
            
            # Estimate output tokens
            output_tokens = len(proofreader.encoder.encode(feedback))
            total_output_tokens += output_tokens
            
            # Save segment results
            segment_result = {
                "segment_id": i,
                "section_indices": segment_indices,
                "section_titles": [sections[idx]["title"] for idx in segment_indices],
                "feedback": feedback,
                "input_tokens": segment_input_tokens,
                "output_tokens": output_tokens
            }
            
            results.append(segment_result)
            
            # Save to files
            with open(os.path.join(output_dir, f"proofreading_segment_{i}.json"), "w") as f:
                json.dump(segment_result, f, indent=2)
            
            with open(os.path.join(output_dir, f"proofreading_segment_{i}.md"), "w") as f:
                f.write(f"# Proofreading Feedback - Segment {i+1}\n\n")
                f.write("## Sections Included\n")
                for idx in segment_indices:
                    f.write(f"- {sections[idx]['title']}\n")
                f.write("\n## Feedback\n\n")
                f.write(feedback)
        
        # Combine results
        combined_path = os.path.join(output_dir, "complete_proofreading_report.md")
        docx_output_path = os.path.join(output_dir, "complete_proofreading_report.docx")
        combined = proofreader.combine_results(results, combined_path, docx_output_path, 
                                              total_input_tokens, total_output_tokens)
        
        print(f"Proofreading complete. Results saved to {output_dir}")
        print(f"Combined report saved to {combined_path}")
        print(f"DOCX report saved to {docx_output_path}")
        
        return {
            "document": document_path,
            "sections": sections,
            "segments": segments,
            "results": results,
            "combined": combined,
            "docx_report": docx_output_path
        }
        
    except Exception as e:
        print(f"Error proofreading document: {str(e)}")
        return {"error": str(e)}


def main():
    """Main entry point with command line argument parsing"""
    parser = argparse.ArgumentParser(description="Document Proofreader using Azure OpenAI")
    parser.add_argument("--input", "-i", help="Input document path")
    parser.add_argument("--output", "-o", help="Output directory")
    parser.add_argument("--config", "-c", action="store_true", help="Show configuration")
    parser.add_argument("--test", "-t", action="store_true", help="Create and proofread test document")
    
    args = parser.parse_args()
    
    # Show configuration if requested
    if args.config:
        try:
            Config.validate()
            Config.display()
            return
        except ValueError as e:
            print(f"Configuration error: {e}")
            return
    
    # Create test document if requested
    if args.test:
        test_doc = create_test_document()
        if not test_doc:
            return
        args.input = test_doc
    
    try:
        # Validate configuration
        Config.validate()
        Config.display()
        
        # Run proofreading
        results = proofread_document(args.input, args.output)
        
        if "error" in results:
            print(f"Error: {results['error']}")
            sys.exit(1)
        else:
            print("\n=== Proofreading completed successfully! ===")
            if "costs" in results:
                costs = results["costs"]
                print(f"Total cost: ${costs['total_cost']:.4f}")
            
    except ValueError as e:
        print(f"Configuration error: {e}")
        print("Please check your .env file contains all required variables:")
        print("- AZURE_OPENAI_ENDPOINT")
        print("- AZURE_OPENAI_API_KEY") 
        print("- AZURE_API_VERSION")
        print("- MODEL_NAME")
        sys.exit(1)
    except Exception as e:
        print(f"Unexpected error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()